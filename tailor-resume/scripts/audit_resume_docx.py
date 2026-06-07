#!/usr/bin/env python3
"""Audit common structural problems in a resume DOCX."""

from __future__ import annotations

import argparse
from collections import Counter
from pathlib import Path
import re
import sys
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.oxml.ns import qn


DATE_RE = re.compile(r"(?:20\d{2}\.\d{2})(?:-(?:20\d{2}\.\d{2}))?")
URL_RE = re.compile(r"https?://\S+")
PLACEHOLDERS = ("xxx", "小标题", "待补充", "姓名", "123456789@qq.com")


def font_tuple(run):
    r_pr = run._element.rPr
    fonts = r_pr.rFonts if r_pr is not None else None
    if fonts is None:
        return (None, None, None)
    return (
        fonts.get(qn("w:eastAsia")),
        fonts.get(qn("w:ascii")),
        fonts.get(qn("w:hAnsi")),
    )


def paragraph_tab_positions(paragraph):
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return []
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        return []
    return [
        (
            tab.get(qn("w:val")),
            tab.get(qn("w:leader")),
            tab.get(qn("w:pos")),
        )
        for tab in tabs
    ]


def has_numbering(paragraph):
    p_pr = paragraph._p.pPr
    return p_pr is not None and p_pr.find(qn("w:numPr")) is not None


def audit(path: Path):
    failures = []
    warnings = []
    facts = []

    try:
        with ZipFile(path) as archive:
            names = set(archive.namelist())
            xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
            comment_parts = [
                name for name in ("word/comments.xml", "word/commentsExtended.xml")
                if name in names
            ]
            if comment_parts:
                failures.append(f"comment parts remain: {', '.join(comment_parts)}")
            media = sorted(name for name in names if name.startswith("word/media/"))
            facts.append(f"embedded media: {len(media)}")
            for placeholder in PLACEHOLDERS:
                if placeholder in xml:
                    failures.append(f"template placeholder remains: {placeholder}")
    except (BadZipFile, KeyError) as exc:
        return [f"invalid DOCX package: {exc}"], warnings, facts

    doc = Document(path)
    fonts = Counter()
    numbered = []
    date_entries = []
    visible_urls = []

    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if has_numbering(paragraph):
            numbered.append((index, text[:50]))
        if URL_RE.search(text):
            visible_urls.extend(URL_RE.findall(text))
        for run in paragraph.runs:
            if run.text.strip():
                fonts[font_tuple(run)] += 1

        if DATE_RE.search(text):
            tabs = paragraph_tab_positions(paragraph)
            date_entries.append((text[:70], "\t" in paragraph.text, tabs))

    if len(fonts) > 1:
        warnings.append(f"multiple explicit font channel sets: {dict(fonts)}")
    elif fonts:
        facts.append(f"font channels: {next(iter(fonts))}")

    if numbered:
        warnings.append(f"paragraph numbering remains: {numbered}")

    tab_positions = Counter()
    for text, has_tab, tabs in date_entries:
        if not has_tab:
            warnings.append(f"date entry has no tab separator: {text}")
        right_tabs = [tab for tab in tabs if tab[0] == "right"]
        if not right_tabs:
            warnings.append(f"date entry has no right-aligned tab stop: {text}")
        for _, leader, pos in right_tabs:
            tab_positions[(leader, pos)] += 1

    if len(tab_positions) > 1:
        warnings.append(f"inconsistent date tab stops: {dict(tab_positions)}")
    elif tab_positions:
        facts.append(f"date tab stop: {next(iter(tab_positions))}")

    facts.append(f"visible URLs: {len(visible_urls)}")
    facts.append(f"body paragraphs: {len(doc.paragraphs)}")
    facts.append(f"tables: {len(doc.tables)}")
    return failures, warnings, facts


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--strict", action="store_true", help="Treat warnings as failure")
    args = parser.parse_args()

    if not args.docx.is_file():
        print(f"FAIL: file not found: {args.docx}")
        return 2

    failures, warnings, facts = audit(args.docx)
    for fact in facts:
        print(f"INFO: {fact}")
    for warning in warnings:
        print(f"WARN: {warning}")
    for failure in failures:
        print(f"FAIL: {failure}")

    if failures or (args.strict and warnings):
        return 1
    print("PASS: structural resume audit completed")
    return 0


if __name__ == "__main__":
    sys.exit(main())
